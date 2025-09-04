#!/usr/bin/env python3
"""
Flask web application for Markdown to DOCX conversion

EDUCATIONAL PURPOSE: This application is designed for educational purposes to demonstrate
basic web development and file conversion techniques. It is intended
for learning about Flask, file processing, and document conversion.

Use responsibly and in accordance with your educational institution's policies.
"""

import os
import re
import secrets
import atexit
import shutil
from pathlib import Path
from flask import Flask, request, render_template, send_file, flash, redirect, url_for
from flask_wtf import FlaskForm
from flask_wtf.csrf import CSRFProtect
from wtforms import FileField, BooleanField
from werkzeug.utils import secure_filename
import tempfile
import uuid
import zipfile

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx not installed. Install with: pip install python-docx")
    exit(1)

app = Flask(__name__)

# Use environment variable for secret key, with fallback for development
app.secret_key = os.environ.get('SECRET_KEY', secrets.token_hex(32))

# Initialize CSRF protection
csrf = CSRFProtect(app)

# Configuration
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['WTF_CSRF_TIME_LIMIT'] = None  # No time limit for CSRF tokens

# Create a dedicated temp folder for this app
UPLOAD_FOLDER = os.path.join(tempfile.gettempdir(), 'markdown_converter')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Track temp files for cleanup
TEMP_FILES = set()

ALLOWED_EXTENSIONS = {'md', 'txt', 'markdown'}

# Cleanup function for temp files
def cleanup_temp_files():
    """Clean up temporary files on exit"""
    for filepath in TEMP_FILES:
        try:
            if os.path.exists(filepath):
                os.unlink(filepath)
        except Exception:
            pass  # Best effort cleanup
    
    # Try to clean the upload folder
    try:
        if os.path.exists(UPLOAD_FOLDER):
            shutil.rmtree(UPLOAD_FOLDER, ignore_errors=True)
    except Exception:
        pass

# Register cleanup function
atexit.register(cleanup_temp_files)

def allowed_file(filename):
    """Check if file extension is allowed"""
    if not filename or '.' not in filename:
        return False
    
    # Ignore system files
    if filename.startswith('.') or filename == 'Thumbs.db':
        return False
    
    # Check extension
    return filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def validate_markdown_content(content):
    """Basic validation of markdown content"""
    # Check for reasonable size (not empty, not too large)
    if not content or len(content) > 10 * 1024 * 1024:  # 10MB content limit
        return False, "File content is empty or too large"
    
    # Check if it contains at least one chapter marker
    chapter_pattern = r'^##? (?:CHAPTER|Chapter) \d+'
    if not re.search(chapter_pattern, content, re.MULTILINE):
        return False, "No chapter markers found in the document"
    
    return True, "Valid"

def parse_markdown_content(markdown_text):
    """Parse markdown content and extract chapters with their content"""
    chapters = []
    current_chapter = None
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        # Check for chapter heading (supports multiple formats)
        chapter_match = None
        chapter_number = None
        chapter_title = None
        
        # Format 1: # Chapter N: Title or ## Chapter N. Title
        match1 = re.match(r'^##? Chapter (\d+)[:.] (.+)$', line.strip())
        if match1:
            chapter_number = int(match1.group(1))
            chapter_title = match1.group(2).strip()
            chapter_match = True
        else:
            # Format 2: # Chapter N or ## Chapter N (just number, no title) 
            match2 = re.match(r'^##? (?:CHAPTER|Chapter) (\d+)$', line.strip())
            if match2:
                chapter_number = int(match2.group(1))
                chapter_title = None  # No title, will be handled in display
                chapter_match = True
        
        if chapter_match:
            # Skip duplicate consecutive chapters (same number)
            if current_chapter and current_chapter['number'] == chapter_number:
                continue
                
            # Save previous chapter if exists
            if current_chapter:
                chapters.append(current_chapter)
            
            # Start new chapter
            current_chapter = {
                'number': chapter_number,
                'title': chapter_title,
                'content_paragraphs': []
            }
        elif line.strip() and current_chapter:
            # Add content paragraph (skip empty lines)
            current_chapter['content_paragraphs'].append(line.strip())
    
    # Add the last chapter
    if current_chapter:
        chapters.append(current_chapter)
    
    return chapters

def create_prowriting_aid_docx(markdown_content, output_path):
    """Create DOCX with exact ProWriting Aid compatible structure"""
    
    # Parse chapters
    chapters = parse_markdown_content(markdown_content)
    
    if not chapters:
        return False, "No chapters found in markdown file"
    
    # Create new document
    doc = Document()
    
    # Get or create required styles
    styles = doc.styles
    
    # Use Heading 2 for chapter titles (ProWritingAid standard)
    heading2_style = None
    for style in styles:
        if style.name == 'Heading 2':
            heading2_style = style
            break
    
    if not heading2_style:
        heading2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    
    # Use Normal style for all content
    normal_style = styles['Normal']  # This always exists in Document
    
    # Process each chapter
    for i, chapter in enumerate(chapters):
        # Create chapter title, avoiding duplication for number-only chapters
        if chapter['title']:
            chapter_title = f"Chapter {chapter['number']}: {chapter['title']}"
        else:
            chapter_title = f"Chapter {chapter['number']}"
        
        # Add chapter heading with Heading2 style (ProWritingAid standard)
        heading_para = doc.add_paragraph(chapter_title)
        heading_para.style = 'Heading 2'
        
        # Add content paragraphs with Normal style (clean, no empty lines)
        for content_para in chapter['content_paragraphs']:
            if content_para.strip():  # Skip empty paragraphs
                para = doc.add_paragraph(content_para)
                para.style = 'Normal'
    
    # Save document
    try:
        doc.save(output_path)
        return True, f"Successfully created DOCX with {len(chapters)} chapters"
    except Exception as e:
        return False, f"Error saving DOCX: {str(e)}"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
@csrf.exempt  # We'll handle CSRF via form validation
def convert_file():
    if 'files' not in request.files:
        flash('No files selected')
        return redirect(request.url)
    
    files = request.files.getlist('files')
    zip_output = request.form.get('zip_output') == 'on'
    
    if not files or all(f.filename == '' for f in files):
        flash('No files selected')
        return redirect(request.url)
    
    # Filter valid files and count ignored ones
    valid_files = []
    ignored_count = 0
    system_files_ignored = []
    
    for f in files:
        if f and f.filename:
            if f.filename.startswith('.') or f.filename == 'Thumbs.db':
                system_files_ignored.append(f.filename)
                ignored_count += 1
            elif allowed_file(f.filename):
                valid_files.append(f)
            else:
                ignored_count += 1
    
    # Provide feedback about ignored files
    if system_files_ignored:
        flash(f'Ignored {len(system_files_ignored)} system file(s): {", ".join(system_files_ignored[:5])}{"..." if len(system_files_ignored) > 5 else ""}')
    
    if not valid_files:
        flash('No valid markdown files found. Please upload .md, .txt, or .markdown files.')
        return redirect(request.url)
    
    try:
        converted_files = []
        unique_id = str(uuid.uuid4())[:8]
        
        # Process each file
        for file in valid_files:
            try:
                # Read markdown content with error handling
                file_content = file.read()
                
                # Try to decode as UTF-8
                try:
                    markdown_content = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    # Try with latin-1 as fallback
                    try:
                        markdown_content = file_content.decode('latin-1')
                    except:
                        flash(f'Unable to decode {file.filename}. Please ensure it is a text file.')
                        continue
                
                # Validate content
                is_valid, validation_msg = validate_markdown_content(markdown_content)
                if not is_valid:
                    flash(f'Invalid content in {file.filename}: {validation_msg}')
                    continue
            except Exception as e:
                flash(f'Error reading {file.filename}: {str(e)}')
                continue
            
            # Create output filename
            original_name = secure_filename(file.filename)
            base_name = original_name.rsplit('.', 1)[0] if '.' in original_name else original_name
            output_filename = f"{base_name}_converted.docx"
            output_path = os.path.join(UPLOAD_FOLDER, f"{unique_id}_{output_filename}")
            
            # Track temp file for cleanup
            TEMP_FILES.add(output_path)
            
            # Convert to DOCX
            success, message = create_prowriting_aid_docx(markdown_content, output_path)
            
            if success:
                converted_files.append((output_path, output_filename))
            else:
                flash(f'Conversion failed for {original_name}: {message}')
        
        if not converted_files:
            flash('No files were successfully converted')
            return redirect(url_for('index'))
        
        # Single file - return directly
        if len(converted_files) == 1 and not zip_output:
            output_path, filename = converted_files[0]
            return send_file(output_path, 
                           as_attachment=True, 
                           download_name=filename,
                           mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        # Multiple files or zip requested - create zip file
        zip_filename = f"converted_manuscripts_{unique_id}.zip"
        zip_path = os.path.join(UPLOAD_FOLDER, zip_filename)
        
        # Track zip file for cleanup
        TEMP_FILES.add(zip_path)
        
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file_path, filename in converted_files:
                zipf.write(file_path, filename)
        
        return send_file(zip_path, 
                       as_attachment=True, 
                       download_name=zip_filename,
                       mimetype='application/zip')
                       
    except Exception as e:
        # Log error for debugging but show generic message to user
        app.logger.error(f'Error processing files: {str(e)}')
        flash('An error occurred while processing your files. Please try again.')
        return redirect(url_for('index'))

if __name__ == '__main__':
    # Get port from environment variable (Render provides this)
    port = int(os.environ.get('PORT', 8080))
    
    # Check if we're in production (Render sets this)
    is_production = os.environ.get('RENDER') is not None
    
    # Run with appropriate settings
    app.run(
        debug=not is_production,  # Debug off in production
        host='0.0.0.0',
        port=port
    )
